from openpyxl import load_workbook
import os
import sys
from docxtpl import DocxTemplate
from petrovich.main import Petrovich
from petrovich.enums import Case, Gender
from docx import settings
from docx import Document
import shutil

document = Document()
document2 = Document()
os.chdir(sys.path[0])
p = Petrovich()

num = ''
place_registration = ''
strData = ''
year_birthday = ''
strUVD = ''
res_strUVD = ''
pathDoc = ''
pathFolder = ''
file_dir = ''
file_dir_docx = ''
firma_pr = ''
files_docx_list = []
directory = []
directory2 = []

# !!!!!!!!!!!!!!!!!!!!!!!!!!! П А Р А М Е Т Р Ы !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!

book = load_workbook('Таблица.xlsx')  # exсel файл с таблицей
doc = DocxTemplate('Шаблон.docx')  # шаблон word документа
name_directory_appeals = 'Обращения'  # название каталога с сохраненными обращениями
sheet = book['4100']  # название книги в exсel файле
line_up = 9  # номер начальной строки в таблице exсel файла
line_down = 42  # номер конечной строки в таблице exсel файла

num_reg_UMVD = '5/1/7750'  # номер исходящего УМВД
num_reg_Sov = '5/1/7751'  # номер исходящего Советский
num_reg_Centr = '5/1/7752'  # номер исходящего Центральный
num_reg_Kriv = '5/1/7753'  # номер исходящего Криволученский
num_reg_Len = '5/1/7754'  # номер исходящего Ленинский
num_reg_Privok = '5/1/7755'  # номер исходящего Привокзальный
num_reg_Zar = '5/1/7756'  # номер исходящего Заречье
num_reg_Skur = '5/1/7757'  # номер исходящего Скуратовский
num_reg_Ilin = '5/1/7758'  # номер исходящего Ильинское
num_reg_Kosog = '5/1/7759'  # номер исходящего Косогорское

date = '12'  # дата месяца исходящего
mou = 'января'  # месяц исходящего
year = '23'  # год исходящего


# !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!


def data_change_format(str_input):
    counter = 0
    str_year = ''
    str_month = ''
    str_day = ''
    for i in str_input:
        if (counter >= 0 and counter < 4):
            str_year += str_input[counter]
        elif (counter >= 5 and counter < 7):
            str_month += str_input[counter]
        elif (counter >= 8 and counter < 10):
            str_day += str_input[counter]
        counter += 1
    str_input = f"{str_day}.{str_month}.{str_year}"
    return str_input


def uvd_writer(str):
    global num
    if (not str.find("ареч") == -1):
        strUVD = "«Зареченский» УМВД России по г. Туле"
        num = num_reg_Zar
        return strUVD
    elif (not str.find("льин") == -1):
        strUVD = "«Ильинское» УМВД России по г. Туле"
        num = num_reg_Ilin
        return strUVD
    elif (not str.find("осог") == -1):
        strUVD = "«Косогорское» УМВД России по г. Туле"
        num = num_reg_Kosog
        return strUVD
    elif (not str.find("волуч") == -1):
        strUVD = "«Криволученский» УМВД России по г. Туле"
        num = num_reg_Kriv
        return strUVD
    elif (not str.find("нинск") == -1):
        strUVD = "«Ленинский» УМВД России по г. Туле"
        num = num_reg_Len
        return strUVD
    elif (not str.find("вокза") == -1):
        strUVD = "«Привокзальный» УМВД России по г. Туле"
        num = num_reg_Privok
        return strUVD
    elif (not str.find("олет") == -1):
        strUVD = "УМВД России по г. Туле"
        num = num_reg_UMVD
        return strUVD
    elif (not str.find("урато") == -1):
        strUVD = "«Скуратовский» УМВД России по г. Туле"
        num = num_reg_Skur
        return strUVD
    elif (not str.find("ветск") == -1):
        strUVD = "«Советский» УМВД России по г. Туле"
        num = num_reg_Sov
        return strUVD
    elif (not str.find("нтра") == -1):
        strUVD = "«Центральный» УМВД России по г. Туле"
        num = num_reg_Centr
        return strUVD


def combine_word_documents(input_files):
    for filnr, file in enumerate(input_files):

        if 'offerte_template' in file:
            file = os.path.join(settings.MEDIA_ROOT, file)

        if filnr == 0:
            merged_document = Document(file)
            merged_document.add_page_break()

        else:
            sub_doc = Document(file)

            if filnr < len(input_files) - 1:
                sub_doc.add_page_break()

            for element in sub_doc.element.body:
                merged_document.element.body.append(element)

    return merged_document


for i in range(line_up, line_down):

    lname = sheet['B' + str(i)].value
    fname = sheet['C' + str(i)].value
    mname = sheet['D' + str(i)].value
    strData = sheet['E' + str(i)].value
    place_registration = sheet['F' + str(i)].value
    notification_date = sheet['H' + str(i)].value
    firma_pr = sheet['L' + str(i)].value
    strData = data_change_format(str(strData))
    year_birthday = strData[6:10]

    notification_date = data_change_format(str(notification_date))

    cased_lastname = p.lastname(lname, Case.GENITIVE, Gender.MALE)
    cased_firstname = p.firstname(fname, Case.GENITIVE, Gender.MALE)
    cased_middlename = p.middlename(mname, Case.GENITIVE, Gender.MALE)
    cased_last_first_middle = f"{cased_lastname} {cased_firstname} {cased_middlename}"

    abbreviated_lastname = f"{lname} {fname[0].upper()}.{mname[0].upper()}."
    cased_abbreviated_lastname = f"{cased_lastname} {fname[0].upper()}.{mname[0].upper()}."

    strSurname = sheet['B' + str(i)].value
    strUVD = uvd_writer(sheet['J' + str(i)].value)

    if firma_pr == 'ф' or firma_pr == 'Ф':
        firma_pr = 'Ф'
    else:
        firma_pr = ' '

    pathDoc = f'{abbreviated_lastname}' + '.docx'

    if not os.path.isdir(name_directory_appeals):
        os.mkdir(name_directory_appeals)

    pathFolder = (sheet['J' + str(i)].value).replace('"', '')
    pathFolder = f"{name_directory_appeals}/{pathFolder}"

    if not os.path.isdir(pathFolder):
        os.mkdir(pathFolder)

    contex = {'strUVD': strUVD,
              'd': date,
              'm': mou,
              'y': year,
              'n': num,
              'cased_last_first_middle': cased_last_first_middle,
              'abbreviated_lastname': abbreviated_lastname,
              'cased_abbreviated_lastname': cased_abbreviated_lastname,
              'strData': strData,
              'place_registration': place_registration,
              'notification_date': notification_date,
              'year_birthday': year_birthday,
              'f': firma_pr
              }
    doc.render(contex)
    doc.save(pathDoc)

    os.replace(f"{pathDoc}", f"{pathFolder}/{pathDoc}")
    print(f"{pathDoc} - создан успешно")
    files_docx_list.append(pathDoc)

document2.save('empty.docx')
file_master = 'empty.docx'

directory = os.listdir(name_directory_appeals)

for file in directory:
    file_dir = os.listdir(f"{name_directory_appeals}/{file}")

    for file_dir_docx in file_dir:
        if len(file_dir) > 1:
            shutil.copy2(f"{name_directory_appeals}/{file}/{file_dir_docx}", f"{file_dir_docx}")
        directory2.append(file_dir_docx)
    if len(directory2) > 1:
        doc2 = combine_word_documents(directory2)
        doc2.save(f"___{file}___.docx")
        for file2 in directory2:
            os.remove(file2)
        shutil.copy2(f"___{file}___.docx", f"{name_directory_appeals}/{file}/___{file}___.docx")
        os.remove(f"___{file}___.docx")
    directory2 = []
os.remove('empty.docx')
